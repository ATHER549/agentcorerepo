import io
import structlog
from docx import Document
from app.extractors.base import BaseExtractor, ExtractionResult

logger = structlog.get_logger()

MAX_PARAGRAPHS = 200


class WordExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        logger.info("Extracting Word content", filename=filename)
        buf = io.BytesIO(file_bytes)
        doc = Document(buf)

        parts = []

        # Extract paragraphs with heading structure
        para_count = 0
        for para in doc.paragraphs:
            if para_count >= MAX_PARAGRAPHS:
                break
            text = para.text.strip()
            if not text:
                continue

            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
            para_count += 1

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                header = rows[0]
                table_str = " | ".join(header) + "\n"
                table_str += " | ".join("---" for _ in header) + "\n"
                for row in rows[1:20]:
                    table_str += " | ".join(row) + "\n"
                parts.append(f"\n### Table {t_idx + 1}\n{table_str}")

        text = "\n\n".join(parts)

        return ExtractionResult(
            text_content=text,
            metadata={"paragraphs_extracted": para_count, "tables_found": len(doc.tables)},
        )
